import docx
from docx.shared import Pt, RGBColor
import re
import os

def update_sds():
    doc_path = r'c:\Users\admin\Downloads\Project_SWP391\Project_SWP391\Project\Template2_SDS Document (1).docx'
    output_path = r'c:\Users\admin\Downloads\Project_SWP391\Project_SWP391\Project\Template2_SDS Document (1)_updated.docx'
    md_path = r'C:\Users\admin\.gemini\antigravity\brain\45799d76-ca1f-4222-bf41-d1ca32b0bd2a\artifacts\sds_section_1_2_packages.md'
    
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f'Error opening docx: {e}')
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        md_lines = f.readlines()
        
    start_idx = -1
    end_idx = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith('1.2 Package Diagram'):
            start_idx = i
        elif text.startswith('1.3 Database Design'):
            end_idx = i
            break
            
    if start_idx == -1 or end_idx == -1:
        print('Could not find section 1.2 or 1.3')
        return
        
    for i in range(end_idx - 1, start_idx, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)
        
    next_p = doc.paragraphs[start_idx + 1] if start_idx + 1 < len(doc.paragraphs) else None
    
    in_table = False
    in_code = False
    code_text = ''
    table_rows = []
    
    for line in md_lines:
        line = line.strip('\n')
        
        if line.startswith('# 1.2 '):
            continue
            
        if line.startswith('```'):
            if in_code:
                in_code = False
                new_p = next_p.insert_paragraph_before(code_text)
                new_p.style.font.name = 'Courier New'
                new_p.style.font.size = Pt(9)
                code_text = ''
            else:
                in_code = True
            continue
            
        if in_code:
            code_text += line + '\n'
            continue
            
        if line.startswith('|'):
            in_table = True
            if line.startswith('| :---') or line.startswith('|---'):
                continue
            
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_rows.append(cells)
            continue
        elif in_table:
            in_table = False
            if table_rows:
                new_p = next_p.insert_paragraph_before('')
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row_data):
                        cell_text = cell_text.replace('`', '').replace('**', '')
                        table.cell(r_idx, c_idx).text = cell_text
                
                p_parent = new_p._element.getparent()
                p_parent.insert(p_parent.index(new_p._element) + 1, table._element)
                p_parent.remove(new_p._element)
                
            table_rows = []
            
        if not line.strip():
            next_p.insert_paragraph_before('')
            continue
            
        new_p = next_p.insert_paragraph_before()
        
        if line.startswith('### '):
            new_p.style = 'Heading 3'
            line = line[4:]
        elif line.startswith('## '):
            new_p.style = 'Heading 2'
            line = line[3:]
        elif line.startswith('- '):
            new_p.style = 'List Paragraph'
            line = line[2:]
            
        parts = re.split(r'(\*\*.*?\*\*|`.*?`)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = new_p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = new_p.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.color.rgb = RGBColor(200, 50, 50)
            else:
                new_p.add_run(part)
                
    # If the last thing was a table, add it
    if in_table and table_rows:
        new_p = next_p.insert_paragraph_before('')
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table.style = 'Table Grid'
        for r_idx, row_data in enumerate(table_rows):
            for c_idx, cell_text in enumerate(row_data):
                cell_text = cell_text.replace('`', '').replace('**', '')
                table.cell(r_idx, c_idx).text = cell_text
        p_parent = new_p._element.getparent()
        p_parent.insert(p_parent.index(new_p._element) + 1, table._element)
        p_parent.remove(new_p._element)
                
    doc.save(output_path)
    print(f'Successfully saved to {output_path}')

if __name__ == '__main__':
    update_sds()

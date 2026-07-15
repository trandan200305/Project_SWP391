import re
import docx

def update_srs():
    doc_path = "Template1_SRS Document (1).docx"
    md_path = r"C:\Users\admin\.gemini\antigravity\brain\45799d76-ca1f-4222-bf41-d1ca32b0bd2a\artifacts\full_use_case_specifications.md"
    output_path = "Template1_SRS Document (1).docx"
    
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Error opening docx: {e}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        md_lines = f.readlines()
        
    # 1. Find section 2 and section 3
    start_idx = -1
    end_idx = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == "2. Use Case Specifications":
            start_idx = i
        elif text == "3. Functional Requirements":
            end_idx = i
            break
            
    if start_idx == -1:
        print("Could not find '2. Use Case Specifications'")
        return
    if end_idx == -1:
        print("Could not find '3. Functional Requirements'")
        return
        
    # 2. Delete all paragraphs between start_idx and end_idx (exclusive)
    # We do this by navigating backwards to not mess up indices
    for i in range(end_idx - 1, start_idx, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)
        
    # 3. Insert markdown content right before section 3
    next_p = doc.paragraphs[end_idx] # This is currently "3. Functional Requirements"
    
    # Skip the first few lines of markdown that repeat "# 2. Use Case Specifications"
    filtered_lines = []
    skip = True
    for line in md_lines:
        line = line.strip('\n')
        if line.startswith('## 2.1'):
            skip = False
        if not skip:
            filtered_lines.append(line)
            
    import re

    for line in filtered_lines:
        if not line.strip():
            next_p.insert_paragraph_before('')
            continue
            
        new_p = next_p.insert_paragraph_before()
        
        # Parse headers
        if line.startswith('### '):
            new_p.style = 'Heading 3'
            line = line[4:]
        elif line.startswith('## '):
            new_p.style = 'Heading 2'
            line = line[3:]
        elif line.startswith('# '):
            new_p.style = 'Heading 1'
            line = line[2:]
        elif line.startswith('- '):
            new_p.style = 'List Paragraph'
            line = line[2:]
            
        # Parse inline bolding with Regex
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = new_p.add_run(part[2:-2])
                run.bold = True
            else:
                new_p.add_run(part)
                
    doc.save(output_path)
    print(f"Successfully saved to {output_path}")

update_srs()

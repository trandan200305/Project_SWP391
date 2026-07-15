import re
import docx
from docx.shared import Pt
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    # Helper to set cell borders if needed, but python-docx default table style Usually has borders.
    pass

def build_srs_tables():
    # Load original template to get a fresh start
    doc_path = "Template1_SRS Document (1)_original.docx" # Wait, do we have the original? 
    # Let's just use the current one and delete everything in section 2
    doc_path = "Template1_SRS Document (1).docx"
    md_path = r"C:\Users\admin\.gemini\antigravity\brain\45799d76-ca1f-4222-bf41-d1ca32b0bd2a\artifacts\full_use_case_specifications.md"
    output_path = "Template1_SRS Document (1).docx"
    
    doc = docx.Document(doc_path)
    
    # 1. Clear Section 2
    start_clear = False
    target_p = None
    for elem in list(doc.element.body):
        if elem.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(elem, doc)
            if p.text.strip() == "2. Use Case Specifications":
                start_clear = True
                continue
            if p.text.strip() == "3. Functional Requirements":
                target_p = p
                break
                
        if start_clear:
            elem.getparent().remove(elem)
            
    # Now all old paragraphs and tables between 2 and 3 are gone.
    # Read Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    ucs = []
    current_uc = None
    current_key = None
    
    for line in lines:
        line = line.strip('\n').strip()
        if not line: continue
        
        if line.startswith('### '):
            if current_uc: ucs.append(current_uc)
            current_uc = {'Title': line[4:], 'Primary Actors':'', 'Secondary Actors':'', 'Description':'', 'Preconditions':'', 'Postconditions':'', 'Normal Sequence/Flow':'', 'Alternative Sequences/Flows':''}
            current_key = None
        elif line.startswith('## '):
            if current_uc:
                ucs.append(current_uc)
                current_uc = None
            # It's a category (e.g. 2.1 Authentication)
            ucs.append({'Category': line[3:]})
        elif current_uc is not None:
            # Check if it's a key
            found_key = False
            for k in ['Primary Actors', 'Secondary Actors', 'Description', 'Preconditions', 'Postconditions', 'Normal Sequence/Flow', 'Alternative Sequences/Flows']:
                if line.startswith(f'**{k}:**'):
                    current_key = k
                    val = line.split(f'**{k}:**', 1)[1].strip()
                    if val:
                        current_uc[current_key] += val + '\n'
                    found_key = True
                    break
            
            if not found_key and current_key:
                current_uc[current_key] += line + '\n'
                
    if current_uc:
        ucs.append(current_uc)
        
    # Helper to insert element before target_p
    def insert_element_before(new_elem, target_paragraph):
        target_paragraph._element.addprevious(new_elem)
        
    # 3. Create elements
    for item in ucs:
        if 'Category' in item:
            p = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
            p.style = 'Heading 2'
            p.add_run(item['Category']).bold = True
            insert_element_before(p._element, target_p)
        else:
            p = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
            p.style = 'Heading 3'
            p.add_run(item['Title']).bold = True
            insert_element_before(p._element, target_p)
            
            # Create Table
            tbl = doc.add_table(rows=6, cols=4)
            tbl.style = 'Table Grid'
            
            # Row 0
            tbl.cell(0,0).text = 'Primary Actors'
            tbl.cell(0,1).text = item['Primary Actors'].strip()
            tbl.cell(0,2).text = 'Secondary Actors'
            tbl.cell(0,3).text = item['Secondary Actors'].strip()
            
            # Merge cols 1 to 3 for the rest
            for r in range(1, 6):
                a = tbl.cell(r, 1)
                b = tbl.cell(r, 3)
                a.merge(b)
                
            tbl.cell(1,0).text = 'Description'
            tbl.cell(1,1).text = item['Description'].strip()
            
            tbl.cell(2,0).text = 'Preconditions'
            tbl.cell(2,1).text = item['Preconditions'].strip()
            
            tbl.cell(3,0).text = 'Postconditions'
            tbl.cell(3,1).text = item['Postconditions'].strip()
            
            tbl.cell(4,0).text = 'Normal Sequence/Flow'
            tbl.cell(4,1).text = item['Normal Sequence/Flow'].strip()
            
            tbl.cell(5,0).text = 'Alternative Sequences/Flows'
            tbl.cell(5,1).text = item['Alternative Sequences/Flows'].strip()
            
            # Add table element before target
            insert_element_before(tbl._element, target_p)
            
            # Add empty paragraph after table for spacing
            spacer = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
            insert_element_before(spacer._element, target_p)

    doc.save(output_path)
    print("Done building tables.")

build_srs_tables()

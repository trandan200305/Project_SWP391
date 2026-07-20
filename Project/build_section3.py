import docx
from docx.oxml import OxmlElement
import re

def insert_element_before(new_elem, target_paragraph):
    target_paragraph._element.addprevious(new_elem)

def build_section3():
    doc_path = "Template1_SRS Document (1).docx"
    md_path = "section3_content.md"
    output_path = "Template1_SRS Document (1).docx"
    
    doc = docx.Document(doc_path)
    
    # 1. Clear Section 3
    start_clear = False
    target_p = None
    for elem in list(doc.element.body):
        if elem.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(elem, doc)
            if p.text.strip() == "3. Functional Requirements":
                start_clear = True
                continue
            if p.text.strip() == "4. Non-Functional Requirements":
                target_p = p
                break
                
        if start_clear:
            elem.getparent().remove(elem)
            
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_data = []
    
    for line in lines:
        line = line.strip('\n').strip()
        if not line:
            continue
            
        if line.startswith('|'):
            if '---|---' in line:
                continue
            cols = [c.strip() for c in line.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_data = [cols]
            else:
                table_data.append(cols)
            continue
            
        # If we just exited a table, render it
        if in_table and not line.startswith('|'):
            in_table = False
            tbl = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            tbl.style = 'Table Grid'
            for r_idx, row in enumerate(table_data):
                for c_idx, val in enumerate(row):
                    tbl.cell(r_idx, c_idx).text = val
            insert_element_before(tbl._element, target_p)
            spacer = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
            insert_element_before(spacer._element, target_p)
            table_data = []

        # Parse text
        p = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
        
        if line.startswith('#### '):
            p.style = 'Heading 4'
            p.add_run(line[5:]).bold = True
        elif line.startswith('### '):
            p.style = 'Heading 3'
            p.add_run(line[4:]).bold = True
        elif line.startswith('## '):
            p.style = 'Heading 2'
            p.add_run(line[3:]).bold = True
        else:
            # Inline bold parsing
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
        insert_element_before(p._element, target_p)
        
    # Flush remaining table
    if in_table:
        tbl = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        tbl.style = 'Table Grid'
        for r_idx, row in enumerate(table_data):
            for c_idx, val in enumerate(row):
                tbl.cell(r_idx, c_idx).text = val
        insert_element_before(tbl._element, target_p)
        spacer = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
        insert_element_before(spacer._element, target_p)

    doc.save(output_path)
    print("Done building section 3.")

build_section3()
